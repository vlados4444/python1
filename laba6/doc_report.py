# doc_report.py
from docx import Document

def create_doc_report(credit_result, deposit_result, file_name="report.docx"):
    doc = Document()
    doc.add_heading("Отчет по банковским услугам", 0)
    
    doc.add_heading("Кредит", level=1)
    doc.add_paragraph(f"Ежемесячный платеж: {credit_result[0]:.2f} руб.")
    doc.add_paragraph(f"Общая сумма выплат: {credit_result[1]:.2f} руб.")
    
    doc.add_heading("Вклад", level=1)
    doc.add_paragraph(f"Общая сумма на счете: {deposit_result[0]:.2f} руб.")
    doc.add_paragraph(f"Начисленные проценты: {deposit_result[1]:.2f} руб.")
    
    doc.save(file_name)