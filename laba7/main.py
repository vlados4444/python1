import tkinter as tk
from tkinter import ttk, messagebox
from bank_services.credit import Credit
from bank_services.installment import Installment
from bank_services.deposit import Deposit
from docx import Document
from openpyxl import Workbook

class BankApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Банковские услуги")

        # Входные поля
        ttk.Label(root, text="Сумма:").grid(row=0, column=0, sticky="w")
        self.amount_entry = ttk.Entry(root)
        self.amount_entry.grid(row=0, column=1)

        ttk.Label(root, text="Срок (мес):").grid(row=1, column=0, sticky="w")
        self.term_entry = ttk.Entry(root)
        self.term_entry.grid(row=1, column=1)

        ttk.Label(root, text="Ставка (%):").grid(row=2, column=0, sticky="w")
        self.rate_entry = ttk.Entry(root)
        self.rate_entry.grid(row=2, column=1)

        ttk.Label(root, text="Тип услуги:").grid(row=3, column=0, sticky="w")
        self.service_var = tk.StringVar()
        self.service_combo = ttk.Combobox(root, textvariable=self.service_var, state="readonly")
        self.service_combo['values'] = ("Кредит", "Рассрочка", "Вклад")
        self.service_combo.grid(row=3, column=1)
        self.service_combo.current(0)

        self.calc_btn = ttk.Button(root, text="Рассчитать", command=self.calculate)
        self.calc_btn.grid(row=4, column=0, columnspan=2, pady=10)

        self.result_text = tk.Text(root, height=8, width=40)
        self.result_text.grid(row=5, column=0, columnspan=2)

        self.save_doc_btn = ttk.Button(root, text="Сохранить в DOCX", command=self.save_doc)
        self.save_doc_btn.grid(row=6, column=0)

        self.save_xls_btn = ttk.Button(root, text="Сохранить в XLSX", command=self.save_xls)
        self.save_xls_btn.grid(row=6, column=1)

        self.calculated_data = []

    def calculate(self):
        try:
            amount = float(self.amount_entry.get())
            term = int(self.term_entry.get())
            rate = float(self.rate_entry.get())
            service = None
            result_str = ""

            if self.service_var.get() == "Кредит":
                service = Credit(amount, term, rate)
                print(service.total_payment)
                payments = service.calculate()
                result_str = f"Ежемесячный платёж: {payments[0]} руб.\nИтого: {service.total_payment} руб."
                self.calculated_data = payments

            elif self.service_var.get() == "Рассрочка":
                service = Installment(amount, term, rate)
                payments = service.calculate()
                result_str = f"Платежей: {term}, каждый: {payments[0]} руб.\nИтого: {service.total_payment} руб."
                self.calculated_data = payments

            elif self.service_var.get() == "Вклад":
                service = Deposit(amount, term, rate)
                future_value = service.calculate()
                result_str = f"Сумма к получению: {future_value} руб."
                self.calculated_data = [future_value]

            self.result_text.delete("1.0", tk.END)
            self.result_text.insert(tk.END, result_str)

        except ValueError:
            messagebox.showerror("Ошибка", "Проверьте правильность ввода данных.")

    def save_doc(self):
        doc = Document()
        doc.add_heading("Отчёт по банковской услуге", 0)
        doc.add_paragraph(self.result_text.get("1.0", tk.END))
        doc.save("report.docx")
        messagebox.showinfo("Успех", "Файл report.docx сохранён.")

    def save_xls(self):
        wb = Workbook()
        ws = wb.active
        ws.title = "Результат"
        ws.append(["Месяц", "Сумма"])
        for i, val in enumerate(self.calculated_data, start=1):
            ws.append([i, val])
        wb.save("report.xlsx")
        messagebox.showinfo("Успех", "Файл report.xlsx сохранён.")


if __name__ == "__main__":
    root = tk.Tk()
    app = BankApp(root)
    root.mainloop()
